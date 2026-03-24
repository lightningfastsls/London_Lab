"""Extract text from parts-finder-plan.docx."""
import io
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document

doc = Document("../parts-finder-plan.docx")
for p in doc.paragraphs:
    print(p.text)

# Also check tables
for i, table in enumerate(doc.tables):
    print(f"\n--- TABLE {i+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
